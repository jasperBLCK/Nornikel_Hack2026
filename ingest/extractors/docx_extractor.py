"""DOCX extractor — headings, paragraphs, tables (best-effort, no LLM)."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document as DocxDocument
from docx.oxml.ns import qn

from ingest.extractors.base import Extractor
from ingest.structures import Page

# DOCX is not natively paginated — we treat the whole document as one
# "page 1" with structured blocks joined by blank lines.
_BLOCK_SEP = "\n\n"
_HEADING_TAG = "heading"


def _block_text(block_elem) -> str:
    return "".join(t.text or "" for t in block_elem.iter(qn("w:t"))).strip()


def _classify_paragraph(p) -> tuple[str, str]:
    text = p.text.strip()
    if not text:
        return ("empty", "")
    style = (p.style.name or "").lower() if p.style else ""
    if style.startswith(_HEADING_TAG):
        level = "".join(ch for ch in style if ch.isdigit()) or "1"
        return (f"h{level}", text)
    if re.match(r"^\d+(\.\d+)*\.?\s+", text):
        return ("numbered", text)
    return ("p", text)


class DocxExtractor(Extractor):
    def extract(self, path: Path) -> list[Page]:
        try:
            doc = DocxDocument(str(path))
        except Exception:  # corrupt or encrypted
            return [Page(page_number=1, text="")]

        blocks: list[str] = []

        # Body paragraphs (with heading detection via style)
        for p in doc.paragraphs:
            kind, text = _classify_paragraph(p)
            if kind == "empty":
                continue
            if kind.startswith("h"):
                blocks.append(f"\n{'#' * int(kind[1:])} {text}\n")
            else:
                blocks.append(text)

        # Tables — render as pipe-separated rows (lossless enough for ETL)
        for t_idx, table in enumerate(doc.tables, 1):
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                blocks.append(f"\n[TABLE {t_idx}]\n" + "\n".join(rows) + "\n[/TABLE]")

        full = _BLOCK_SEP.join(b for b in blocks if b)
        return [Page(page_number=1, text=full)]